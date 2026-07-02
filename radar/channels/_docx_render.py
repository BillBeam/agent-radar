"""Render a digest's full markdown (`Digest.markdown`) → a .docx (bytes), preserving the
4-axis structure so it's a real deep-read document, not a formatless blob.

Mapping (structure fidelity is the whole point):
  # / ## / ###        → Heading 1 / 2 / 3          (title / sections / per-item)
  a whole `**line**`  → Heading 4                   (the ①②③④ axis sub-heads — must stay visible)
  inline **b** / *i*  → bold / italic runs
  [text](url)         → real docx hyperlink (blue underline)
  - / *  item         → List Bullet
  > quote             → Quote style
  blank / ---         → skipped (paragraph spacing carries it)

Content is NEVER regenerated — this only re-formats the already-rendered `Digest.markdown`.
python-docx is imported here (only loaded when the docx delivery path actually runs).
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# one token = bold | italic | [text](url)
_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*].*?\*|\[[^\]]+\]\([^)]+\))")
_LINK = re.compile(r"^\[([^\]]+)\]\(([^)]+)\)$")
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BOLD_LINE = re.compile(r"^\*\*(.+)\*\*$")
_BULLET = re.compile(r"^[-*]\s+(.*)$")
_QUOTE = re.compile(r"^>\s?(.*)$")


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """python-docx has no native hyperlink — build the w:hyperlink element by hand (blue+underline)."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rpr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _add_runs(paragraph, text: str) -> None:
    """Parse inline **bold** / *italic* / [text](url) and append the corresponding runs."""
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            paragraph.add_run(tok[2:-2]).bold = True
        elif (m := _LINK.match(tok)):
            _add_hyperlink(paragraph, m.group(2), m.group(1))
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            paragraph.add_run(tok[1:-1]).italic = True
        else:
            paragraph.add_run(tok)


def markdown_to_docx(md: str) -> bytes:
    """Digest.markdown → .docx bytes (structure preserved). Pure — no network, no regeneration."""
    doc = Document()
    for raw in md.split("\n"):
        line = raw.rstrip()
        if not line.strip() or line.strip() == "---":
            continue
        if (m := _HEADING.match(line)):
            _add_runs(doc.add_heading(level=min(len(m.group(1)), 4)), m.group(2))
        elif (m := _BOLD_LINE.match(line)):                 # ①②③④ axis sub-heads → Heading 4
            _add_runs(doc.add_heading(level=4), m.group(1))
        elif (m := _BULLET.match(line)):
            _add_runs(doc.add_paragraph(style="List Bullet"), m.group(1))
        elif (m := _QUOTE.match(line)):
            _add_runs(doc.add_paragraph(style="Quote"), m.group(1))
        else:
            _add_runs(doc.add_paragraph(), line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
