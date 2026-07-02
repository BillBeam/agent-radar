"""md → docx renderer — structure fidelity (pure, no network)."""
from __future__ import annotations

import io

from docx import Document

from radar.channels._docx_render import markdown_to_docx

SAMPLE = """# Agent Radar · 2026-06-30
> 扫描 28 源 · 候选 117
## 🆕 今日新增
### [1] [Some Paper Title](https://arxiv.org/abs/2606.1)
*arXiv (agent/LLM)*　paper · eval
**① 核心机制**
这是**核心**内容，含 *斜体* 与 [内链](https://x.org/a)。
- 第一点
- 第二点
**takeaway**
一句话总结。
"""


def _doc(md: str) -> Document:
    return Document(io.BytesIO(markdown_to_docx(md)))


def test_heading_hierarchy_and_axis_subheads():
    styles = [p.style.name for p in _doc(SAMPLE).paragraphs]
    assert "Heading 1" in styles           # # 标题
    assert "Heading 2" in styles           # ## 今日新增
    assert "Heading 3" in styles           # ### [1] ...
    assert styles.count("Heading 4") == 2  # **① 核心机制** + **takeaway** → 四轴小标题保真
    assert styles.count("List Bullet") == 2
    assert "Quote" in styles               # > 扫描…


def test_bold_line_becomes_clean_heading():
    doc = _doc("**② 证据 / 数据**")
    h = next(p for p in doc.paragraphs if p.style.name == "Heading 4")
    assert h.text == "② 证据 / 数据"        # ** stripped, not double-wrapped


def test_inline_bold_and_italic_runs():
    p = _doc("这是**粗**和*斜*收尾。").paragraphs[0]
    assert any(r.bold for r in p.runs)
    assert any(r.italic for r in p.runs)


def test_hyperlink_element_present():
    doc = _doc("见 [某文](https://arxiv.org/abs/2606.9) 原文。")
    assert "hyperlink" in doc.paragraphs[0]._p.xml         # w:hyperlink element built
    assert any("arxiv.org/abs/2606.9" in getattr(r, "target_ref", "")
               for r in doc.part.rels.values())            # url lives in the relationship


def test_empty_and_hr_lines_skipped():
    doc = _doc("# H\n\n---\n正文")
    texts = [p.text for p in doc.paragraphs]
    assert "H" in texts and "正文" in texts and "---" not in texts
